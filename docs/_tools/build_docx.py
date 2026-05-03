from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "docs" / "Dossier_Documentation_GameConnect_BTS_SIO_SLAM.docx"


ACCENT = "1F4E79"
LIGHT = "EAF3F8"
DARK = "17324D"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor.from_string(color)
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, title in enumerate(headers):
        set_cell_text(hdr[i], title, bold=True, color="FFFFFF")
        shade(hdr[i], ACCENT)
        if widths:
            hdr[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(DARK if level == 1 else ACCENT)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, LIGHT)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)
    p.add_run(f"\n{body}")
    doc.add_paragraph()


def section_break(doc):
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Aptos Display"
    styles["Heading 1"].font.size = Pt(20)
    styles["Heading 2"].font.size = Pt(15)
    styles["Heading 3"].font.size = Pt(12)
    return doc


def build():
    doc = setup_document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GameConnect")
    run.bold = True
    run.font.size = Pt(30)
    run.font.color.rgb = RGBColor.from_string(DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("Dossier de documentation - BTS SIO SLAM")
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor.from_string(ACCENT)

    doc.add_paragraph()
    add_note(
        doc,
        "Objet du dossier",
        "Documentation d'examen couvrant le contexte, l'analyse, la conception, la réalisation, les tests, la documentation technique/utilisateur et le bilan du projet GameConnect.",
    )
    add_table(
        doc,
        ["Élément", "Valeur"],
        [
            ["Projet", "Plateforme sociale e-sport avec matching de joueurs"],
            ["Formation", "BTS SIO option SLAM"],
            ["Frontend", "React 18, Vite, Axios, Tailwind CSS"],
            ["Backend", "FastAPI, Python, Pydantic, JWT, bcrypt"],
            ["Base de données", "MySQL 8"],
            ["Déploiement", "Docker Compose, Nginx"],
            ["Dépôt Git", "https://github.com/Farounaga/Esportapp.git"],
        ],
        [4.5, 12.5],
    )

    add_heading(doc, "Sommaire", 1)
    add_numbered(
        doc,
        [
            "Contexte et cahier des charges",
            "Analyse",
            "Conception",
            "Réalisation",
            "Tests",
            "Documentation technique et utilisateur",
            "Bilan",
        ],
    )

    section_break(doc)
    add_heading(doc, "01 - Contexte et cahier des charges", 1)
    doc.add_paragraph(
        "GameConnect est une application web destinée aux joueurs souhaitant trouver des coéquipiers compatibles selon leurs jeux, leur niveau, leur région, leur fuseau horaire et leurs objectifs."
    )
    add_heading(doc, "Objectifs fonctionnels", 2)
    add_bullets(
        doc,
        [
            "Créer un compte, se connecter et gérer un profil de joueur.",
            "Ajouter des jeux avec niveau, rang, temps de jeu et favori.",
            "Calculer des suggestions de coéquipiers avec un score de compatibilité.",
            "Accepter ou refuser un match.",
            "Échanger par messagerie uniquement entre matchs acceptés.",
            "Consulter statistiques, recherche et notifications.",
        ],
    )
    add_heading(doc, "User stories principales", 2)
    add_table(
        doc,
        ["ID", "User story", "Priorité"],
        [
            ["US01", "En tant que visiteur, je veux créer un compte afin d'utiliser la plateforme.", "Haute"],
            ["US02", "En tant qu'utilisateur, je veux me connecter afin d'accéder à mon espace.", "Haute"],
            ["US03", "En tant que joueur, je veux modifier mon profil afin de présenter mon niveau.", "Haute"],
            ["US04", "En tant que joueur, je veux ajouter mes jeux afin d'améliorer le matching.", "Haute"],
            ["US05", "En tant que joueur, je veux obtenir des suggestions de coéquipiers.", "Haute"],
            ["US06", "En tant que joueur, je veux envoyer des messages à mes matchs acceptés.", "Moyenne"],
        ],
        [2, 11.5, 3],
    )

    section_break(doc)
    add_heading(doc, "02 - Analyse", 1)
    add_heading(doc, "Acteurs et cas d'utilisation", 2)
    add_table(
        doc,
        ["Acteur", "Cas d'utilisation"],
        [
            ["Visiteur", "Consulter l'accueil, s'inscrire, se connecter"],
            ["Joueur connecté", "Gérer profil, jeux, matchs, messages, notifications et statistiques"],
            ["Système", "Vérifier JWT, calculer matching, enregistrer activité"],
            ["Administrateur technique", "Installer, configurer, initialiser la BDD et déployer"],
        ],
        [4, 12],
    )
    add_heading(doc, "Règles métier", 2)
    add_table(
        doc,
        ["Code", "Règle"],
        [
            ["RM01", "L'email et le pseudo sont uniques."],
            ["RM02", "Le mot de passe contient au moins 8 caractères, une lettre et un chiffre."],
            ["RM03", "Les mots de passe sont hachés avec bcrypt."],
            ["RM04", "Les routes protégées nécessitent un JWT valide."],
            ["RM05", "Un utilisateur ne peut pas ajouter deux fois le même jeu."],
            ["RM06", "Un profil privé ne doit pas être proposé dans le matching."],
            ["RM07", "Un message n'est possible qu'entre deux joueurs avec match accepté."],
            ["RM08", "La suppression de message est logique avec `deleted_at`."],
        ],
        [2.2, 14.3],
    )
    add_heading(doc, "Choix technologiques", 2)
    add_table(
        doc,
        ["Technologie", "Justification"],
        [
            ["React + Vite", "SPA moderne, composants réutilisables, build rapide."],
            ["FastAPI", "API REST performante avec validation Pydantic et Swagger automatique."],
            ["MySQL", "SGBDR adapté au MCD/MLD et aux relations du projet."],
            ["JWT + bcrypt", "Authentification stateless et stockage sécurisé des mots de passe."],
            ["Docker Compose", "Déploiement reproductible en services db, api et frontend."],
        ],
        [4, 12],
    )

    section_break(doc)
    add_heading(doc, "03 - Conception", 1)
    add_heading(doc, "MCD et MLD", 2)
    doc.add_paragraph(
        "Le MCD est centré sur l'entité Utilisateur, liée à son profil, ses jeux, ses préférences, ses matchs, messages, notifications et traces d'activité."
    )
    add_table(
        doc,
        ["Table", "Rôle", "Relations principales"],
        [
            ["users", "Compte et authentification", "Référence centrale"],
            ["user_profiles", "Profil social et confidentialité", "1-1 avec users"],
            ["games", "Catalogue de jeux", "N-N via user_games"],
            ["user_games", "Jeux déclarés par les utilisateurs", "Unique user_id + game_id"],
            ["matches", "Mises en relation", "Deux FK vers users"],
            ["messages", "Messagerie privée", "sender_id et receiver_id vers users"],
            ["notifications", "Notifications utilisateur", "user_id et from_user_id"],
            ["user_activity_logs", "Historique d'activité", "FK vers users"],
        ],
        [3.5, 5.5, 7],
    )
    add_heading(doc, "Architecture applicative", 2)
    add_bullets(
        doc,
        [
            "Vue : pages et composants React dans `frontend/src`.",
            "Contrôleurs : routes FastAPI dans `API/app/routes`.",
            "Services métier : authentification, matching et activité dans `API/app/services`.",
            "Persistance : MySQL via `DatabaseSession`.",
            "Communication : JSON HTTP avec Axios et token Bearer.",
        ],
    )
    add_heading(doc, "Séquences principales", 2)
    add_table(
        doc,
        ["Séquence", "Résumé"],
        [
            ["Inscription", "Le frontend envoie les données, l'API valide, hash le mot de passe, crée users/profil et retourne un JWT."],
            ["Matching", "L'API vérifie le JWT, récupère profil et jeux, calcule un score pondéré, crée des matchs pending."],
            ["Messagerie", "L'API vérifie l'existence d'un match accepted avant d'insérer le message."],
        ],
        [3.5, 13],
    )

    section_break(doc)
    add_heading(doc, "04 - Réalisation", 1)
    add_heading(doc, "Environnement", 2)
    add_table(
        doc,
        ["Élément", "Fichier ou commande"],
        [
            ["Backend", "`API/app/main.py`, `uvicorn app.main:app --reload --port 8000`"],
            ["Frontend", "`frontend/src/App.jsx`, `npm run dev`"],
            ["Base locale", "`localsetup/database.sql`"],
            ["Base Docker", "`deploy/docker/mysql-init/00_schema.sql`"],
            ["Déploiement", "`docker compose up -d --build`"],
        ],
        [4, 12],
    )
    add_heading(doc, "Endpoints principaux", 2)
    add_table(
        doc,
        ["Domaine", "Endpoints"],
        [
            ["Auth", "POST /register, POST /login"],
            ["Profil", "GET /profile, PUT /profile, GET /user/activity-stats"],
            ["Jeux", "GET /games, GET/POST/PUT/DELETE /user/games"],
            ["Matching", "POST /matches, GET /matches, accept, reject"],
            ["Messages", "GET /messages, GET /messages/{id}, POST /messages, DELETE /messages/{id}"],
            ["Notifications", "GET /notifications, unread-count, read, read-all, delete"],
            ["Stats/Recherches", "GET /stats/*, GET /search/*"],
        ],
        [3.5, 13],
    )

    section_break(doc)
    add_heading(doc, "05 - Tests", 1)
    add_heading(doc, "Plan de tests", 2)
    add_table(
        doc,
        ["ID", "Scénario", "Résultat attendu"],
        [
            ["TF01", "Inscription", "Compte, profil et JWT créés."],
            ["TF02", "Connexion", "JWT retourné avec identifiants valides."],
            ["TF03", "Ajout d'un jeu", "Association créée sans doublon."],
            ["TF04", "Matching", "Suggestions ou message informatif."],
            ["TF05", "Message sans match accepté", "HTTP 403."],
            ["TF06", "Déploiement Docker", "Services healthy."],
        ],
        [2, 6, 8],
    )
    add_heading(doc, "Tests disponibles", 2)
    add_bullets(
        doc,
        [
            "`frontend/src/__tests__` : tests Jest pour AuthContext, Avatar, FormInput, Messages, ThemeContext et Toast.",
            "`API/tests/test_api.py` : script d'intégration API avec requests.",
            "Commandes : `npm test`, `npm run test:coverage`, `python tests/test_api.py`.",
        ],
    )
    add_note(
        doc,
        "Rapport",
        "Avant la soutenance, compléter le tableau de résultats dans docs/05_tests/README.md après exécution des commandes sur l'environnement final.",
    )

    section_break(doc)
    add_heading(doc, "06 - Documentation technique et utilisateur", 1)
    add_heading(doc, "Installation Docker", 2)
    add_numbered(
        doc,
        [
            "Cloner le dépôt GitHub.",
            "Configurer les variables d'environnement et les secrets.",
            "Lancer `docker compose up -d --build`.",
            "Vérifier `docker compose ps`.",
            "Accéder au frontend sur http://localhost:8080 et à Swagger sur http://localhost:8000/docs.",
        ],
    )
    add_heading(doc, "Guide utilisateur", 2)
    add_table(
        doc,
        ["Action", "Étapes"],
        [
            ["Créer un compte", "Page inscription, email, pseudo, mot de passe, profil initial."],
            ["Compléter son profil", "Page Profil, saisir région, bio, niveau et préférences."],
            ["Ajouter un jeu", "Page Jeux, sélectionner un jeu, niveau, rang, heures et favori."],
            ["Trouver un coéquipier", "Page Matching, lancer la recherche, accepter ou refuser."],
            ["Envoyer un message", "Page Messages, choisir un match accepté et envoyer le texte."],
        ],
        [4, 12],
    )
    add_heading(doc, "Dépannage", 2)
    add_table(
        doc,
        ["Problème", "Solution"],
        [
            ["Erreur CORS", "Vérifier `CORS_ORIGINS`."],
            ["API ne contacte pas MySQL", "Vérifier `DB_HOST`, `DB_USER`, `DB_PASS`, `DB_NAME`."],
            ["Aucun match trouvé", "Ajouter des jeux et profils de test compatibles."],
            ["Token refusé", "Se reconnecter et vérifier `JWT_SECRET`."],
        ],
        [5, 11],
    )

    section_break(doc)
    add_heading(doc, "07 - Bilan", 1)
    add_heading(doc, "Points forts", 2)
    add_bullets(
        doc,
        [
            "Architecture claire et séparée entre React, FastAPI et MySQL.",
            "API REST documentée avec Swagger.",
            "Sécurité de base : JWT, bcrypt, validation Pydantic.",
            "Déploiement Docker reproductible.",
            "Tests frontend présents et script de test API.",
        ],
    )
    add_heading(doc, "Axes d'amélioration", 2)
    add_table(
        doc,
        ["Axe", "Priorité"],
        [
            ["Transformer le script API en suite Pytest automatisée.", "Haute"],
            ["Ajouter une CI GitHub Actions.", "Haute"],
            ["Ajouter notifications temps réel avec WebSocket ou SSE.", "Moyenne"],
            ["Ajouter modération, signalement et blocage utilisateur.", "Haute"],
            ["Renforcer RGPD : export et suppression de compte.", "Haute"],
            ["Ajouter observabilité : logs structurés et monitoring.", "Moyenne"],
        ],
        [12, 3],
    )
    add_note(
        doc,
        "Conclusion",
        "GameConnect constitue une base solide pour un projet SLAM : conception de données, API REST, frontend React, tests, sécurité et déploiement sont couverts. Les prochaines étapes concernent surtout l'automatisation des tests, la modération et la conformité RGPD.",
    )

    doc.core_properties.title = "Dossier de documentation GameConnect - BTS SIO SLAM"
    doc.core_properties.subject = "Documentation projet"
    doc.core_properties.author = "Equipe GameConnect"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
