from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT_DIR = ROOT / "docs" / "livrables_exam" / "word"

ACCENT = "1F4E79"
DARK = "17324D"
LIGHT = "EAF3F8"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text, bold=False, color=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True, color="FFFFFF")
        shade(table.rows[0].cells[i], ACCENT)
        if widths:
            table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
            if widths:
                cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_title(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor.from_string(DARK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor.from_string(ACCENT)
    doc.add_paragraph()


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor.from_string(DARK if level == 1 else ACCENT)
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_note(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, LIGHT)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK)
    p.add_run(f"\n{body}")
    doc.add_paragraph()


def new_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].paragraph_format.space_after = Pt(6)
    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        styles[style_name].font.name = "Aptos Display"
    add_title(doc, title, subtitle)
    return doc


def save(doc, filename, title):
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.core_properties.title = title
    doc.core_properties.subject = "Livrable examen BTS SIO SLAM"
    doc.core_properties.author = "Equipe GameConnect"
    path = OUT_DIR / filename
    doc.save(path)
    return path


def suivi_trello():
    doc = new_doc("Suivi de projet Trello", "GameConnect - BTS SIO SLAM")
    add_heading(doc, "Objectif", 1)
    doc.add_paragraph(
        "Le suivi de projet permet de visualiser l'avancement des taches, les priorites et la progression du projet GameConnect."
    )
    add_heading(doc, "Lien de suivi", 1)
    doc.add_paragraph("https://trello.com/b/QJRSOkjn/agile-scrum-gameconnect")
    add_heading(doc, "Organisation du tableau", 1)
    add_table(
        doc,
        ["Colonne", "Utilite"],
        [
            ["Backlog", "Idees, besoins et user stories non commencees."],
            ["A faire", "Taches selectionnees pour l'iteration en cours."],
            ["En cours", "Taches en developpement ou en redaction."],
            ["Tests / validation", "Taches terminees techniquement mais a verifier."],
            ["Termine", "Taches validees et integrees."],
        ],
        [4, 12],
    )
    add_heading(doc, "Exemples de cartes", 1)
    add_table(
        doc,
        ["Carte", "Type", "Priorite"],
        [
            ["Creation du compte utilisateur", "Fonctionnalite", "Haute"],
            ["Connexion JWT", "Fonctionnalite", "Haute"],
            ["Ajout d'un jeu au profil", "Fonctionnalite", "Haute"],
            ["Calcul du score de matching", "Fonctionnalite", "Haute"],
            ["Messagerie entre matchs acceptes", "Fonctionnalite", "Moyenne"],
            ["Schema MySQL", "Base de donnees", "Haute"],
            ["Tests frontend Jest", "Test", "Moyenne"],
            ["Deploiement Docker Compose", "Deploiement", "Haute"],
        ],
        [8, 4, 3],
    )
    add_note(
        doc,
        "Utilisation a l'examen",
        "Ce livrable montre la methode de travail, la repartition des taches, les priorites et l'avancement du projet.",
    )
    return save(doc, "Suivi_de_projet_Trello_GameConnect.docx", "Suivi de projet Trello GameConnect")


def note_cadrage():
    doc = new_doc("Note de cadrage", "GameConnect - BTS SIO SLAM")
    add_heading(doc, "Contexte", 1)
    doc.add_paragraph(
        "GameConnect est une plateforme sociale e-sport permettant a des joueurs de creer un profil, declarer leurs jeux et trouver des coequipiers compatibles."
    )
    add_heading(doc, "Probleme identifie", 1)
    doc.add_paragraph(
        "Les joueurs utilisent souvent plusieurs outils differents pour trouver des partenaires. Cette dispersion rend la recherche de coequipiers moins efficace."
    )
    add_heading(doc, "Objectif du projet", 1)
    doc.add_paragraph(
        "Centraliser les informations importantes d'un joueur et proposer un matching base sur les jeux communs, le niveau, la region, le fuseau horaire et les objectifs."
    )
    add_heading(doc, "Perimetre", 1)
    add_table(
        doc,
        ["Inclus", "Hors perimetre initial"],
        [
            ["Inscription, connexion, profil", "Paiement ou abonnement"],
            ["Catalogue et jeux utilisateur", "Application mobile native"],
            ["Matching, messages, notifications", "Chat temps reel WebSocket"],
            ["Statistiques, recherche, Docker", "Connexion Steam, Riot ou Twitch"],
        ],
        [8, 8],
    )
    add_heading(doc, "Contraintes", 1)
    add_bullets(
        doc,
        [
            "Projet web compatible BTS SIO SLAM.",
            "Base de donnees relationnelle MySQL.",
            "API REST separee du frontend.",
            "Authentification securisee.",
            "Documentation claire pour installation, tests et deploiement.",
        ],
    )
    return save(doc, "Note_de_Cadrage_GameConnect.docx", "Note de cadrage GameConnect")


def cahier_specifications():
    doc = new_doc("Cahier des specifications", "GameConnect - BTS SIO SLAM")
    add_heading(doc, "Description generale", 1)
    doc.add_paragraph(
        "GameConnect est une application web composee d'un frontend React, d'une API FastAPI et d'une base MySQL. Elle propose une experience complete autour du profil e-sport et de la mise en relation."
    )
    add_heading(doc, "Exigences fonctionnelles", 1)
    add_table(
        doc,
        ["ID", "Exigence", "Priorite"],
        [
            ["F01", "Un visiteur peut creer un compte.", "Haute"],
            ["F02", "Un utilisateur peut se connecter.", "Haute"],
            ["F03", "Un utilisateur peut consulter et modifier son profil.", "Haute"],
            ["F04", "Un utilisateur peut ajouter, modifier ou supprimer ses jeux.", "Haute"],
            ["F05", "Le systeme propose des matchs compatibles.", "Haute"],
            ["F06", "Un utilisateur peut accepter ou refuser un match.", "Haute"],
            ["F07", "Un utilisateur peut envoyer des messages a un match accepte.", "Moyenne"],
            ["F08", "Un utilisateur peut consulter ses notifications.", "Moyenne"],
        ],
        [2, 11, 3],
    )
    add_heading(doc, "Exigences non fonctionnelles", 1)
    add_table(
        doc,
        ["ID", "Exigence", "Solution retenue"],
        [
            ["NF01", "Securiser les mots de passe.", "bcrypt"],
            ["NF02", "Proteger les routes privees.", "JWT Bearer"],
            ["NF03", "Valider les donnees.", "Pydantic"],
            ["NF04", "Persister les donnees.", "MySQL 8"],
            ["NF05", "Faciliter le deploiement.", "Docker Compose"],
            ["NF06", "Documenter l'API.", "Swagger FastAPI"],
        ],
        [2, 8, 6],
    )
    add_heading(doc, "Criteres de validation", 1)
    add_bullets(
        doc,
        [
            "L'application demarre en local ou via Docker.",
            "L'inscription et la connexion fonctionnent.",
            "Un utilisateur peut completer son profil et ajouter des jeux.",
            "Le matching retourne une liste ou un message explicite.",
            "Les routes protegees refusent les requetes sans JWT.",
        ],
    )
    return save(doc, "Cahier_des_specifications_GameConnect.docx", "Cahier des specifications GameConnect")


def plan_agile():
    doc = new_doc("Plan de projet agile", "GameConnect - BTS SIO SLAM")
    add_heading(doc, "Methode", 1)
    doc.add_paragraph(
        "Le projet peut etre presente avec une methode agile simplifiee de type Scrum/Kanban. Les taches sont suivies dans Trello et decoupees en user stories."
    )
    add_heading(doc, "Roles", 1)
    add_table(
        doc,
        ["Role", "Responsabilite"],
        [
            ["Product owner", "Priorise les besoins et valide les fonctionnalites."],
            ["Developpeur backend", "API FastAPI, securite, base de donnees."],
            ["Developpeur frontend", "Interfaces React, services Axios, tests UI."],
            ["Testeur", "Scenarios de tests et validation fonctionnelle."],
            ["DevOps", "Docker, variables d'environnement et deploiement."],
        ],
        [5, 11],
    )
    add_heading(doc, "Backlog par lots", 1)
    add_table(
        doc,
        ["Lot", "Contenu", "Priorite"],
        [
            ["Lot 1", "Authentification, compte utilisateur, JWT.", "Haute"],
            ["Lot 2", "Profil joueur et catalogue de jeux.", "Haute"],
            ["Lot 3", "Matching et acceptation/refus.", "Haute"],
            ["Lot 4", "Messagerie et notifications.", "Moyenne"],
            ["Lot 5", "Statistiques, recherche et activite.", "Moyenne"],
            ["Lot 6", "Docker, tests et documentation.", "Haute"],
        ],
        [3, 10, 3],
    )
    add_heading(doc, "Definition of Done", 1)
    add_bullets(
        doc,
        [
            "Le code est integre au projet.",
            "Les donnees sont sauvegardees correctement.",
            "Les erreurs principales sont gerees.",
            "La fonctionnalite est testee.",
            "La documentation est mise a jour si necessaire.",
        ],
    )
    return save(doc, "Plan_de_Projet_Agile_GameConnect.docx", "Plan de projet agile GameConnect")


def stack_technique():
    doc = new_doc("Stack technique", "GameConnect - BTS SIO SLAM")
    add_heading(doc, "Vue d'ensemble", 1)
    doc.add_paragraph("L'application est separee en trois services : frontend React, API FastAPI et base de donnees MySQL.")
    add_heading(doc, "Frontend", 1)
    add_table(
        doc,
        ["Technologie", "Role"],
        [
            ["React 18", "Construction de l'interface utilisateur."],
            ["Vite", "Serveur de developpement et outil de build."],
            ["React Router", "Navigation entre les pages."],
            ["Axios", "Appels HTTP vers l'API."],
            ["Tailwind CSS", "Mise en forme responsive."],
            ["Jest / Testing Library", "Tests frontend."],
        ],
        [5, 11],
    )
    add_heading(doc, "Backend et base de donnees", 1)
    add_table(
        doc,
        ["Technologie", "Role"],
        [
            ["FastAPI", "Framework API REST."],
            ["Pydantic", "Validation des donnees."],
            ["PyJWT", "Creation et verification des tokens JWT."],
            ["bcrypt", "Hachage des mots de passe."],
            ["MySQL 8", "SGBDR relationnel."],
            ["Docker Compose", "Deploiement reproductible."],
        ],
        [5, 11],
    )
    add_note(
        doc,
        "Justification",
        "Cette stack montre une architecture web moderne, une API REST securisee, une base relationnelle exploitable pour MCD/MLD et un deploiement reproductible.",
    )
    return save(doc, "Stack_Technique_GameConnect.docx", "Stack technique GameConnect")


def explication_mcd_mld():
    doc = new_doc("Explication MCD / MLD", "GameConnect - BTS SIO SLAM")
    add_heading(doc, "Objectif du modele", 1)
    doc.add_paragraph(
        "Le modele represente les utilisateurs, leur profil de joueur, les jeux pratiques, les mises en relation, les messages, les notifications et l'activite."
    )
    add_heading(doc, "Entites principales", 1)
    add_table(
        doc,
        ["Entite", "Description"],
        [
            ["Utilisateur", "Compte avec email, pseudo et mot de passe hache."],
            ["Profil", "Informations sociales et e-sport de l'utilisateur."],
            ["Jeu", "Catalogue des jeux disponibles."],
            ["UtilisateurJeu", "Association entre un utilisateur et un jeu."],
            ["Match", "Mise en relation entre deux utilisateurs."],
            ["Message", "Message prive entre deux utilisateurs."],
            ["Notification", "Information envoyee a un utilisateur."],
            ["Activite", "Trace technique d'une action utilisateur."],
        ],
        [5, 11],
    )
    add_heading(doc, "Cardinalites", 1)
    add_table(
        doc,
        ["Relation", "Cardinalite", "Justification"],
        [
            ["Utilisateur - Profil", "1,1", "Un utilisateur possede un seul profil."],
            ["Utilisateur - Jeu", "N,N", "Un joueur peut jouer a plusieurs jeux."],
            ["Utilisateur - Match", "1,N", "Un joueur peut avoir plusieurs matchs."],
            ["Utilisateur - Message", "1,N", "Un joueur peut envoyer et recevoir plusieurs messages."],
            ["Utilisateur - Notification", "1,N", "Un joueur peut recevoir plusieurs notifications."],
        ],
        [5, 3, 8],
    )
    add_heading(doc, "Passage au MLD", 1)
    add_bullets(
        doc,
        [
            "Ajout de cles primaires `id`.",
            "Ajout de cles etrangeres vers `users` et `games`.",
            "Creation de `user_games` pour la relation N-N.",
            "Contrainte unique `(user_id, game_id)`.",
            "Index SQL pour ameliorer les recherches.",
        ],
    )
    add_note(
        doc,
        "Fichiers associes",
        "MCD Mermaid : docs/03_conception/mcd.mmd ; MLD Mermaid : docs/03_conception/mld.mmd ; Schema SQL : deploy/docker/mysql-init/00_schema.sql.",
    )
    return save(doc, "Explication_MCD_MLD_GameConnect.docx", "Explication MCD MLD GameConnect")


def main():
    paths = [
        suivi_trello(),
        note_cadrage(),
        cahier_specifications(),
        plan_agile(),
        stack_technique(),
        explication_mcd_mld(),
    ]
    for path in paths:
        print(path)


if __name__ == "__main__":
    main()
